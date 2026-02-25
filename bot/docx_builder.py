"""DOCX generator: assembles a GOST 7.32-2017 formatted document from typed blocks."""

from __future__ import annotations

import io
from datetime import datetime

from docx import Document
from docx.shared import Pt, Cm, Mm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from bot.converter import (
    Block, HeadingBlock, ParagraphBlock, ListBlock, CodeBlock, TableBlock,
    ImageBlock, BlockquoteBlock,
    Run, Metadata,
)
from bot.styles import (
    FONT_NAME, FONT_SIZE, CODE_FONT_NAME, CODE_FONT_SIZE, TABLE_FONT_SIZE,
    LINE_SPACING, PARAGRAPH_INDENT, WORK_TYPES,
    setup_page, setup_default_style, set_run_font, set_paragraph_format,
    DEFAULT_UNIVERSITY, DEFAULT_INSTITUTE, DEFAULT_DEPARTMENT,
    DEFAULT_CITY, DEFAULT_GROUP,
)


def build_docx(blocks: list[Block], metadata: Metadata, work_type: str = "lab") -> io.BytesIO:
    """Build a GOST-formatted DOCX from parsed blocks and metadata. Returns a BytesIO."""
    doc = Document()
    setup_page(doc)
    setup_default_style(doc)

    _add_title_page(doc, metadata, work_type)
    _add_page_break(doc)
    _add_toc(doc, blocks)
    _add_page_break(doc)
    _render_blocks(doc, blocks)
    _add_page_numbers(doc)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# --- Helpers ---

def _centered_line(doc, text, bold=True, font_size=Pt(12), space_before=Pt(0), space_after=Pt(0)):
    """Add a centered paragraph with given formatting."""
    p = doc.add_paragraph()
    set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                         first_line_indent=Cm(0), line_spacing=1.0,
                         space_before=space_before, space_after=space_after)
    if text:
        run = p.add_run(text)
        set_run_font(run, font_size=font_size, bold=bold)
    return p


def _right_block_line(doc, text, bold=False, font_size=Pt(12), left_indent=Cm(10)):
    """Add a left-aligned paragraph with a big left indent (appears on the right half)."""
    p = doc.add_paragraph()
    set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                         first_line_indent=Cm(0), line_spacing=1.0,
                         space_before=Pt(0), space_after=Pt(0))
    p.paragraph_format.left_indent = left_indent
    if text:
        run = p.add_run(text)
        set_run_font(run, font_size=font_size, bold=bold)
    return p


# --- Title page (matching sample format) ---

def _add_title_page(doc: Document, meta: Metadata, work_type: str) -> None:
    """Generate a university-style cover page matching the GUMRF sample."""
    university = meta.university or DEFAULT_UNIVERSITY
    institute = meta.institute or DEFAULT_INSTITUTE
    department = meta.department or DEFAULT_DEPARTMENT
    title = meta.title or ""
    author = meta.author or ""
    group = meta.group or DEFAULT_GROUP
    teacher = meta.teacher or ""
    year = meta.year or str(datetime.now().year)
    city = meta.city or DEFAULT_CITY
    work_label = WORK_TYPES.get(work_type, WORK_TYPES["lab"])
    work_number = meta.work_number or ""
    subject = meta.subject or ""

    sz = Pt(12)

    # University name block — centered, bold
    _centered_line(doc, university, bold=True, font_size=sz)

    # Separator line
    _centered_line(doc, "–" * 56, bold=True, font_size=sz)

    # Institute
    _centered_line(doc, institute, bold=True, font_size=sz)

    # Empty line
    _centered_line(doc, "", font_size=sz)

    # Department
    _centered_line(doc, department, bold=True, font_size=sz)

    # Spacing
    for _ in range(2):
        _centered_line(doc, "", font_size=sz)

    # "ОТЧЕТ"
    _centered_line(doc, "ОТЧЕТ", bold=True, font_size=Pt(16))

    # "по лабораторной работе №X"
    work_line = f"по {work_label}"
    if work_number:
        work_line += f" №{work_number}"
    _centered_line(doc, work_line, bold=True, font_size=Pt(14))

    # Subject line
    if subject:
        _centered_line(doc, f"по дисциплине «{subject}»", bold=True, font_size=sz)

    # Title
    if title:
        _centered_line(doc, f"«{title}»", bold=True, font_size=sz,
                        space_before=Pt(6))

    # Spacing before author block
    for _ in range(3):
        _centered_line(doc, "", font_size=sz)

    left = Cm(10)

    # "Выполнил: студент группы XX"
    _right_block_line(doc, f"Выполнил: студент группы {group}", bold=True, font_size=sz, left_indent=left)

    # Empty line
    _right_block_line(doc, "", font_size=sz, left_indent=left)

    # Author name + signature line
    if author:
        _right_block_line(doc, f"{author}   ______________", font_size=sz, left_indent=left)
    else:
        _right_block_line(doc, "________________________________________   ______________", font_size=sz, left_indent=left)

    # Labels
    p = _right_block_line(doc, "", font_size=Pt(10), left_indent=left)
    run = p.add_run("          (фамилия, имя, отчество)                    (подпись)")
    set_run_font(run, font_size=Pt(10))

    # Spacing
    _right_block_line(doc, "", font_size=sz, left_indent=left)

    # "Руководитель:"
    _right_block_line(doc, "Руководитель:", bold=True, font_size=sz, left_indent=left)

    # Empty line
    _right_block_line(doc, "", font_size=sz, left_indent=left)

    # Teacher name + signature line
    if teacher:
        _right_block_line(doc, f"{teacher}   ______________", font_size=sz, left_indent=left)
    else:
        _right_block_line(doc, "________________________________________   ______________", font_size=sz, left_indent=left)

    # Labels
    p = _right_block_line(doc, "", font_size=Pt(10), left_indent=left)
    run = p.add_run("          (фамилия, имя, отчество)                    (подпись)")
    set_run_font(run, font_size=Pt(10))

    # Spacing
    _right_block_line(doc, "", font_size=sz, left_indent=left)

    # "Представлена на кафедру:" line
    _right_block_line(doc, f"Представлена на кафедру:   «___» _____________ {year} г.",
                      bold=True, font_size=sz, left_indent=left)

    # Spacing before city/year
    _centered_line(doc, "", font_size=sz)

    # City
    _centered_line(doc, city, bold=True, font_size=sz)

    # Year
    _centered_line(doc, year, bold=True, font_size=sz)


# --- TOC (generated from actual headings) ---

def _add_toc(doc: Document, blocks: list[Block]) -> None:
    """Generate a real table of contents from heading blocks."""
    # Title
    _centered_line(doc, "СОДЕРЖАНИЕ", bold=True, font_size=Pt(14),
                   space_after=Pt(12))

    # Collect headings
    headings = [(b.level, b.text) for b in blocks if isinstance(b, HeadingBlock)]
    if not headings:
        p = doc.add_paragraph()
        set_paragraph_format(p, first_line_indent=Cm(0))
        run = p.add_run("(Содержание пусто — нет заголовков)")
        set_run_font(run, font_size=Pt(12), italic=True)
        return

    # Section numbering counters
    counters = [0, 0, 0]  # H1, H2, H3

    for level, text in headings:
        if level == 1:
            counters[0] += 1
            counters[1] = 0
            counters[2] = 0
            num = str(counters[0])
        elif level == 2:
            counters[1] += 1
            counters[2] = 0
            num = f"{counters[0]}.{counters[1]}"
        else:
            counters[2] += 1
            num = f"{counters[0]}.{counters[1]}.{counters[2]}"

        indent = Cm((level - 1) * 1.0)

        p = doc.add_paragraph()
        set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                             first_line_indent=Cm(0),
                             space_before=Pt(2), space_after=Pt(2),
                             line_spacing=1.5)
        p.paragraph_format.left_indent = indent

        # Use tab leader dots
        display_text = text.upper() if level == 1 else text
        entry = f"{num} {display_text}"
        run = p.add_run(entry)
        font_sz = Pt(14)
        set_run_font(run, font_size=font_sz, bold=(level == 1))

        # Add dotted tab stop + page placeholder
        _add_toc_tab_stop(p)
        tab_run = p.add_run("\t")
        set_run_font(tab_run, font_size=font_sz)


def _add_toc_tab_stop(paragraph) -> None:
    """Add a right-aligned tab stop with dot leader at the right margin."""
    pPr = paragraph._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:leader"), "dot")
    # Right margin position (~16.5cm from left margin for A4 with 30mm+15mm margins)
    tab.set(qn("w:pos"), "9356")  # in twips (16.5cm)
    tabs.append(tab)
    pPr.append(tabs)


# --- Page break ---

def _add_page_break(doc: Document) -> None:
    """Add a page break."""
    doc.add_page_break()


# --- Block rendering ---

def _render_blocks(doc: Document, blocks: list[Block], number_headings: bool = True) -> None:
    """Render all blocks into the document."""
    counters = [0, 0, 0]  # H1, H2, H3
    for block in blocks:
        if isinstance(block, HeadingBlock):
            section_number = ""
            if number_headings:
                if block.level == 1:
                    counters[0] += 1
                    counters[1] = 0
                    counters[2] = 0
                    section_number = str(counters[0])
                elif block.level == 2:
                    counters[1] += 1
                    counters[2] = 0
                    section_number = f"{counters[0]}.{counters[1]}"
                else:
                    counters[2] += 1
                    section_number = f"{counters[0]}.{counters[1]}.{counters[2]}"
            _render_heading(doc, block, section_number=section_number)
        elif isinstance(block, ParagraphBlock):
            _render_paragraph(doc, block)
        elif isinstance(block, ListBlock):
            _render_list(doc, block)
        elif isinstance(block, CodeBlock):
            _render_code_block(doc, block)
        elif isinstance(block, TableBlock):
            _render_table(doc, block)
        elif isinstance(block, ImageBlock):
            _render_image(doc, block)
        elif isinstance(block, BlockquoteBlock):
            _render_blockquote(doc, block)


def _render_heading(doc: Document, block: HeadingBlock, section_number: str = "") -> None:
    """Render heading with GOST formatting. H1: centered, bold, uppercase. H2+: bold, indented."""
    p = doc.add_paragraph()
    prefix = f"{section_number} " if section_number else ""
    if block.level == 1:
        set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                             first_line_indent=Cm(0),
                             space_before=Pt(24), space_after=Pt(12))
        run = p.add_run(f"{prefix}{block.text.upper()}")
        set_run_font(run, font_size=Pt(14), bold=True)
        _set_heading_outline_level(p, 0)
    elif block.level == 2:
        set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                             first_line_indent=PARAGRAPH_INDENT,
                             space_before=Pt(18), space_after=Pt(6))
        run = p.add_run(f"{prefix}{block.text}")
        set_run_font(run, font_size=Pt(14), bold=True)
        _set_heading_outline_level(p, 1)
    else:
        set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                             first_line_indent=PARAGRAPH_INDENT,
                             space_before=Pt(12), space_after=Pt(6))
        run = p.add_run(f"{prefix}{block.text}")
        set_run_font(run, font_size=Pt(14), bold=True)
        _set_heading_outline_level(p, min(block.level - 1, 8))


def _set_heading_outline_level(paragraph, level: int) -> None:
    """Set the outline level on a paragraph so it appears in the TOC."""
    pPr = paragraph._p.get_or_add_pPr()
    outline = OxmlElement("w:outlineLvl")
    outline.set(qn("w:val"), str(level))
    pPr.append(outline)


def _render_paragraph(doc: Document, block: ParagraphBlock) -> None:
    """Render a paragraph with formatted runs."""
    p = doc.add_paragraph()
    set_paragraph_format(p)
    for run_data in block.runs:
        run = p.add_run(run_data.text)
        if run_data.code:
            set_run_font(run, font_name=CODE_FONT_NAME, font_size=CODE_FONT_SIZE)
        else:
            set_run_font(run, bold=run_data.bold, italic=run_data.italic)


def _render_list(doc: Document, block: ListBlock) -> None:
    """Render a list (ordered or unordered) with proper indentation."""
    for i, item_runs in enumerate(block.items):
        p = doc.add_paragraph()
        set_paragraph_format(p, first_line_indent=Cm(0))
        p.paragraph_format.left_indent = PARAGRAPH_INDENT

        if block.ordered:
            prefix = f"{i + 1}. "
        else:
            prefix = "– "
        run = p.add_run(prefix)
        set_run_font(run)

        for run_data in item_runs:
            run = p.add_run(run_data.text)
            if run_data.code:
                set_run_font(run, font_name=CODE_FONT_NAME, font_size=CODE_FONT_SIZE)
            else:
                set_run_font(run, bold=run_data.bold, italic=run_data.italic)


def _render_code_block(doc: Document, block: CodeBlock) -> None:
    """Render a code block with monospace font and border."""
    # Language label
    if block.language:
        lbl = doc.add_paragraph()
        set_paragraph_format(lbl, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                             first_line_indent=Cm(0),
                             space_before=Pt(6), space_after=Pt(0),
                             line_spacing=1.0)
        lbl.paragraph_format.left_indent = PARAGRAPH_INDENT
        r = lbl.add_run(f"Листинг ({block.language}):")
        set_run_font(r, font_size=Pt(10), italic=True)

    p = doc.add_paragraph()
    set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                         first_line_indent=Cm(0),
                         space_before=Pt(3) if block.language else Pt(6),
                         space_after=Pt(6),
                         line_spacing=1.0)
    p.paragraph_format.left_indent = PARAGRAPH_INDENT
    _add_paragraph_border(p)
    _add_paragraph_shading(p, "F2F2F2")

    # Split code into lines and use add_break() for proper line breaks
    lines = block.code.split("\n")
    for i, line in enumerate(lines):
        run = p.add_run(line)
        set_run_font(run, font_name=CODE_FONT_NAME, font_size=CODE_FONT_SIZE)
        if i < len(lines) - 1:
            run.add_break()


def _add_paragraph_border(paragraph) -> None:
    """Add a thin border around a paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "4")
        border.set(qn("w:color"), "808080")
        pBdr.append(border)
    pPr.append(pBdr)


def _add_paragraph_shading(paragraph, color: str) -> None:
    """Add background shading to a paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color)
    pPr.append(shd)


def _render_table(doc: Document, block: TableBlock) -> None:
    """Render a table with borders and centered alignment."""
    num_cols = len(block.headers) if block.headers else (len(block.rows[0]) if block.rows else 0)
    if num_cols == 0:
        return

    num_rows = (1 if block.headers else 0) + len(block.rows)
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_borders(table)

    row_idx = 0
    if block.headers:
        for col_idx, header in enumerate(block.headers):
            cell = table.cell(0, col_idx)
            p = cell.paragraphs[0]
            set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=Cm(0))
            run = p.add_run(header)
            set_run_font(run, font_size=TABLE_FONT_SIZE, bold=True)
        row_idx = 1

    for row_data in block.rows:
        for col_idx, cell_text in enumerate(row_data):
            if col_idx >= num_cols:
                break
            cell = table.cell(row_idx, col_idx)
            p = cell.paragraphs[0]
            set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=Cm(0))
            run = p.add_run(cell_text)
            set_run_font(run, font_size=TABLE_FONT_SIZE)
        row_idx += 1


def _render_image(doc: Document, block: ImageBlock) -> None:
    """Render an image block. Downloads from URL if available."""
    image_data = _download_image(block.url)
    if image_data:
        doc.add_picture(image_data, width=Cm(14))
        last_para = doc.paragraphs[-1]
        last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p = doc.add_paragraph()
        set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=Cm(0))
        run = p.add_run(f"[Изображение: {block.alt or block.url}]")
        set_run_font(run, italic=True)
    if block.alt:
        p = doc.add_paragraph()
        set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                             first_line_indent=Cm(0),
                             space_before=Pt(2), space_after=Pt(6))
        run = p.add_run(block.alt)
        set_run_font(run, font_size=Pt(12), italic=True)


def _download_image(url: str) -> io.BytesIO | None:
    """Try to download an image from URL. Returns BytesIO or None on failure."""
    if not url or not url.startswith(("http://", "https://")):
        return None
    try:
        import urllib.request
        req = urllib.request.Request(url, headers={"User-Agent": "Mozilla/5.0"})
        with urllib.request.urlopen(req, timeout=10) as resp:
            data = resp.read(5 * 1024 * 1024)  # max 5MB
        return io.BytesIO(data)
    except Exception:
        return None


def _render_blockquote(doc: Document, block: BlockquoteBlock) -> None:
    """Render a blockquote with increased indent and italic."""
    p = doc.add_paragraph()
    set_paragraph_format(p, first_line_indent=Cm(0),
                         space_before=Pt(6), space_after=Pt(6))
    p.paragraph_format.left_indent = Cm(2)
    for run_data in block.runs:
        run = p.add_run(run_data.text)
        set_run_font(run, italic=True, bold=run_data.bold)


def _set_table_borders(table) -> None:
    """Apply borders to all cells in a table."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    borders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "000000")
        borders.append(border)
    tblPr.append(borders)


# --- Page numbering ---

def _add_page_numbers(doc: Document) -> None:
    """Add centered page numbers in the footer, skipping the first page."""
    section = doc.sections[0]
    section.different_first_page_header_footer = True

    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = p.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    run._r.append(fldChar1)

    run2 = p.add_run()
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = " PAGE "
    run2._r.append(instrText)

    run3 = p.add_run()
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")
    run3._r.append(fldChar2)

    run4 = p.add_run("1")
    set_run_font(run4, font_size=FONT_SIZE)

    run5 = p.add_run()
    fldChar3 = OxmlElement("w:fldChar")
    fldChar3.set(qn("w:fldCharType"), "end")
    run5._r.append(fldChar3)

    first_footer = section.first_page_footer
    first_footer.is_linked_to_previous = False
